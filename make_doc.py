#!/usr/bin/env python3
"""Generate the P4/BMv2 experiment log docx for the thesis."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name='微软雅黑', size=10.5, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # East Asian font
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), name)


def heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=(0x1F, 0x3B, 0x5C))
    elif level == 2:
        set_font(run, size=13, bold=True, color=(0x2C, 0x5F, 0x8A))
    elif level == 3:
        set_font(run, size=11.5, bold=True)
    else:
        set_font(run, size=10.5, bold=True)
    p.paragraph_format.space_before = Pt(10 if level <= 2 else 6)
    p.paragraph_format.space_after = Pt(4)
    return p


def body(doc, text, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(2)
    return p


def bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=size)
    p.paragraph_format.space_after = Pt(1)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, name='Consolas', size=9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(4)
    # light gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p


def table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        set_font(run, size=9.5, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = t.cell(i + 1, j)
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=9.5)
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    return t


doc = Document()
# default font
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ============ 封面 ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('P4/BMv2 可编程交换机实验记录')
set_font(run, size=22, bold=True)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('双节点通讯 → 16节点流量控制 → 分层强化学习全局调度')
set_font(run, size=13, color=(0x60, 0x60, 0x60))
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('用于论文编写 · 持续跟进更新')
set_font(run, size=11, color=(0x90, 0x90, 0x90))
doc.add_paragraph()

# ============ 0. 环境配置 ============
heading(doc, '0. 实验环境配置', 1)
body(doc, '本实验在 Windows 11 + Docker Desktop 上运行,通过 p4lang/p4app 容器提供 Linux 环境。')
table(doc,
      ['项目', '配置'],
      [
          ['宿主系统', 'Windows 11 Home China (22H3 / 10.0.22631)'],
          ['Docker', 'Docker Desktop 29.4.3'],
          ['容器', 'p4app(镜像 p4lang/p4app,最新版)'],
          ['容器入口', '--privileged --entrypoint bash(默认 p4apprunner 不接受 -c)'],
          ['挂载', 'E:\\p4-workspace -> /workspace'],
          ['P4 编译器', 'p4c 1.2.0(v1model 架构)'],
          ['数据面', 'BMv2 simple_switch(单线程软件交换机)'],
          ['控制面', 'simple_switch_CLI(Thrift, 端口 9090)'],
          ['脚本语言', 'Python 2.7(容器内),脚本需 python2 运行'],
          ['网络技术', 'Linux 网络命名空间(netns)+ veth pair + 静态 ARP'],
      ],
      col_widths=[1.5, 4.3])

heading(doc, '0.1 容器重建与启动命令', 2)
code_block(doc, 'docker run -d --privileged --entrypoint bash \\\n'
               '  -v "E:\\p4-workspace:/workspace" --name p4app \\\n'
               '  p4lang/p4app:latest -c "while true; do sleep 3600; done"')
code_block(doc, 'docker start p4app\n'
               'docker exec p4app bash -c "cd /workspace && python2 -u run_global_demo.py"')
body(doc, '注意:容器内只有 Python 2.7 工具链(simple_switch_CLI 是 python2 脚本),脚本必须用 python2 运行;'
          '不能用宿主机的 Python 3。')

heading(doc, '0.2 目录结构(整理后)', 2)
code_block(doc, 'E:\\p4-workspace\\\n'
               '├── exp2_node\\             # 实验一:双节点通讯\n'
               '├── exp16_tokenbucket\\     # 实验二:16节点 + 令牌桶流量控制\n'
               '└── rl_global_flow_control\\ # 实验三:分层RL全局调度 + DCQCN + Credit')

doc.add_page_break()

# ============ 实验一 ============
heading(doc, '1. 实验一:双节点通讯(2-node)', 1)
heading(doc, '1.1 目标', 2)
body(doc, '在 BMv2 可编程交换机上实现两个网络命名空间主机之间的 IPv4 转发,验证 P4 数据面转发通路,'
          '并建立后续实验的基础设施。')

heading(doc, '1.2 文件与步骤', 2)
table(doc,
      ['文件', '作用'],
      [
          ['simple_router.p4', 'P4 程序:IPv4 LPM 最长前缀匹配转发'],
          ['run_dual_node.py', '一键脚本:编译→拓扑→BMv2→流表→ping→抓包'],
          ['topo.py', 'Mininet 版本拓扑(备用)'],
          ['capture_demo.html', 'ICMP 抓包可视化页面'],
      ],
      col_widths=[1.6, 4.2])
heading(doc, '1.3 每步做了什么', 3)
bullet(doc, '① 编译:P4 程序用 p4c 编译成 BMv2 的 JSON(build/simple_router.json)')
bullet(doc, '② 建拓扑:创建 ns-h1(10.0.0.1)和 ns-h2(10.0.0.2)两个命名空间,各一根 veth 网线连到交换机端口 1、2')
bullet(doc, '③ 静态 ARP:两台主机互相绑对方的真实 MAC(nud permanent),避免 ARP 广播干扰')
bullet(doc, '④ 启动 BMv2:simple_switch 加载 JSON,端口 9090 提供 Thrift 控制接口')
bullet(doc, '⑤ 下流表:10.0.0.1/32 → 端口1 + MAC1;10.0.0.2/32 → 端口2 + MAC2')
bullet(doc, '⑥ 测试:ns-h1 ping ns-h2,验证双向通信')
bullet(doc, '⑦ 抓包:tcpdump 在交换机两端口抓包,合并成 ping_capture.pcap,用 Wireshark 可视化')

heading(doc, '1.4 结果与关键调试经验', 2)
table(doc,
      ['现象', '根因', '修复'],
      [
          ['h1→h2 0% 丢包;h2→h1 首次 66% 丢包', 'BMv2 启动时序抖动', '加长启动等待 sleep'],
          ['UDP/TCP 全不通,只有 ping 通', 'BMv2 转发改 TTL 后不重算 UDP 校验和,接收端 InCsumErrors 静默丢弃', 'P4 forward action 里 h.udp.checksum = 0'],
          ['跨命名空间通信失败', 'ARP 广播不被交换机转发', 'ip neigh ... nud permanent 静态绑定真实 MAC'],
          ['流表用假 MAC 导致不通', '接收端内核不认', '流表必须填主机真实 MAC'],
      ],
      col_widths=[2.2, 2.6, 1.9])

doc.add_page_break()

# ============ 实验二 ============
heading(doc, '2. 实验二:16 节点 + 令牌桶流量控制(16-node token bucket)', 1)
heading(doc, '2.1 目标', 2)
body(doc, '把双节点扩展到 16 节点,并用 P4 v1model meter(令牌桶)实现逐目的流量控制:'
          '对指定目的主机限速,超速的包被丢弃,不限速的保持全通。')

heading(doc, '2.2 文件与步骤', 2)
table(doc,
      ['文件', '作用'],
      [
          ['simple_router16.p4', 'P4:16端口 LPM 转发 + meter(16, MeterType.bytes) 逐目的限速'],
          ['run_16node.py', '一键脚本:编译→16命名空间→BMv2 16口→静态ARP→120条路由→限速演示'],
          ['traffic_control_demo.html', '结果可视化(拓扑 + 令牌桶动画)'],
          ['traffic_demo_beginners.html', '小白版可视化(快递站比喻)'],
      ],
      col_widths=[1.9, 3.9])
heading(doc, '2.3 流量控制原理', 2)
body(doc, 'meter 是 2-rate 3-color 令牌桶:每个目的主机一个桶(destIdx = 目的IP末字节 - 1)。'
          '桶按 rate 补充令牌,包过闸要消耗令牌;GREEN(0)转发,非 GREEN 丢弃。'
          'rate 管持续速率(单位 bytes/µs),burst 管突发容忍度。')
code_block(doc, 'meter(16, MeterType.bytes) m_dst;\n'
               'm_dst.execute_meter<bit<32>>(m.destIdx, color);\n'
               'if (color == 0) { ipv4_lpm.apply(); } else { drop(); }')
heading(doc, '2.4 每步做了什么', 3)
bullet(doc, '① 编译 simple_router16.p4')
bullet(doc, '② 建 16 个命名空间 ns-h1..16(10.0.0.1..16/24),各一根 veth 连交换机端口 1..16')
bullet(doc, '③ 全对全静态 ARP(120 条,静默执行避免刷屏)')
bullet(doc, '④ 启动 BMv2 16 端口(无 --log-console,日志写文件避免死锁)')
bullet(doc, '⑤ 下 16 条 LPM 路由(真实 MAC)')
bullet(doc, '⑥ 连通性检查:120 对 ping 全通')
bullet(doc, '⑦ 流量控制:meter_set_rates 给 h2 限速 8Mbps、h16 不限速;打 15.7Mbps 流对比')
heading(doc, '2.5 结果与关键调试经验', 2)
table(doc,
      ['项目', '结果/经验'],
      [
          ['连通性', '120/120 双向全通'],
          ['限速效果', 'h2(8Mbps)300发→150到(50%丢);h16(不限速)300→300(0%丢)'],
          ['burst 必须 ≥ 包大小', 'burst=800 < 1400B 时全丢;burst≥1500 时低速率流正常通过'],
          ['速率换算', 'BMv2 meter rate 单位 bytes/µs,从 bps 需 ÷8÷1e6(少÷8 会把 8M 配成 64M)'],
          ['--log-console', '逐包日志拖垮吞吐,去掉;输出重定向到文件而非 PIPE(否则缓冲填满死锁)'],
      ],
      col_widths=[1.7, 4.1])

doc.add_page_break()

# ============ 实验三 ============
heading(doc, '3. 实验三:分层强化学习全局调度 + DCQCN + Credit(当前)', 1)
heading(doc, '3.1 目标', 2)
body(doc, '在 16 节点自由双向通讯基础上,用软件侧分层强化学习(RL)替代数据面令牌桶做全局流量控制,'
          '结合 DCQCN 拥塞控制(ECN 信号→量化降速)和 Credit 信用流控(防接收缓冲丢包)。'
          '节点间自由随机通讯(每节点 1~3 个目的),并预留 MPTCP 多路径扩展抽象。')

heading(doc, '3.2 文件与步骤', 2)
table(doc,
      ['文件', '作用'],
      [
          ['simple_router_global.p4', 'P4:转发 + egress 用 deq_timedelta 判拥塞打 ECN 标记 + 计数寄存器(无 meter)'],
          ['flow.py', '数据模型:Flow(id,src,dst,subflows[]) / Subflow(rate,path,rtt,ecn);生成随机连接图;MPTCP 预留'],
          ['rl_train.py', '分层 Q-learning:高层 Q_tree(选流)+ 低层 Q_flow(速率),交替冻结训练,导出 policy.json(含环境指纹)'],
          ['global_scheduler.py', '控制面调度器:高层选激活流→低层定速率→DCQCN量化降速→pacing'],
          ['credit_flow.py', '主机侧 credit:接收方授信、发送方额度内发'],
          ['run_global_demo.py', '主脚本:编译→拓扑→随机连接图→超发→调度→丢包归因→写 live_stats.json'],
          ['monitor.html', '实时监控页:轮询 live_stats.json 渲染 ECN/活跃流/速率/丢包'],
          ['LOSS_ATTRIBUTION.md', '丢包归因表(按位置/约束分类所有遇过的丢包)'],
          ['README.md', '实验说明'],
      ],
      col_widths=[1.9, 3.9])

heading(doc, '3.3 总体架构', 2)
code_block(doc, '┌───────────── 软件侧 ─────────────┐\n'
               '│ ① 分层 RL 全局调度器(控制面)     │\n'
               '│    高层 Q_tree: 本轮激活哪些流    │\n'
               '│    低层 Q_flow: 给选中流定速率    │\n'
               '│    交替冻结训练(论文 Algorithm 1)│\n'
               '│    策略持久化 + 环境指纹检测      │\n'
               '│ ② DCQCN: ECN信号→量化降速→缓恢复 │\n'
               '│ ③ Credit: 接收方授信,发送方额度内发│\n'
               '└─────────────────────────────────┘\n'
               '               ↓ 下发速率/额度\n'
               '┌───────────── 数据面 ─────────────┐\n'
               '│ P4: 16口转发 + egress 排队等待    │\n'
               '│ (deq_timedelta)超阈值→打ECN标记    │\n'
               '└─────────────────────────────────┘')

heading(doc, '3.4 每步做了什么', 3)
bullet(doc, '① 编译 simple_router_global.p4(转发 + ECN)')
bullet(doc, '② 建 16 命名空间 + veth,BMv2 16 口,LPM 路由,静态 ARP')
bullet(doc, '③ 加载/训练分层策略:有 policy.json 且指纹匹配就复用,否则 rl_train.py 重训')
bullet(doc, '④ 生成自由连接图:flow.py 每节点随机 1~3 个目的,形成有向流(A→B 和 B→A 独立,双向)')
bullet(doc, '⑤ Phase A(超发):所有流高速发送制造拥塞,触发 ECN')
bullet(doc, '⑥ Phase B(调度):调度器分层决策(选流+速率)+ DCQCN,协调各流')
bullet(doc, '⑦ 丢包归因:统计收发、丢包率,检查接收端校验和/缓冲约束')
bullet(doc, '⑧ 实时监控:每轮写 live_stats.json,monitor.html 实时刷新')

heading(doc, '3.5 分层 RL 的设计(借鉴论文)', 2)
body(doc, '借鉴 "AllReduce Scheduling with Hierarchical DRL"(Wei et al.),把调度拆成两层 Q 表,同时保留原实验的 Q-learning 特性:')
table(doc,
      ['层', '状态', '动作', '奖励'],
      [
          ['高层 Q_tree(选流)', '拥塞等级(0/1/2) × 剩余工作占比', '激活 all / half / quarter', '吞吐 + 阶段完成奖励(+5)'],
          ['低层 Q_flow(速率)', '拥塞等级(0/1/2)', '速率倍率 ×1.2/1.0/0.7/0.4', '吞吐 − 0.4·时延 − 1.5·丢包'],
      ],
      col_widths=[1.7, 1.9, 1.7, 1.9])
body(doc, '训练交替冻结:冻结 Q_flow 训 Q_tree,再冻结 Q_tree 训 Q_flow(论文 Algorithm 1)。'
          '策略持久化:policy.json 存两层策略 + Q 表 + 环境指纹 + 连接图指纹,环境没变就复用,变了才重训。')

heading(doc, '3.5.1 增量探索机制(基础模型 + 新连接图微调)', 3)
body(doc, '面对随机连接图每次可能不同的特点,引入"基础模型 + 增量探索"范式,用两类指纹判定:')
table(doc,
      ['判定条件', '动作', '说明'],
      [
          ['环境指纹 + 连接图指纹都匹配', '直接复用', '不训练,策略完全可用'],
          ['连接图变了、训练参数没变', '增量微调', '在旧 Q 表上小学习率(0.05)继续训,保留基础策略快速适配新结构'],
          ['无旧策略 / 训练参数变了', '从零重训', '首次或参数变更时全量训练'],
      ],
      col_widths=[2.4, 1.2, 2.4])
body(doc, '连接图指纹 = 当前 (src,dst) 连接集合的哈希。这样不同 seed 生成的不同随机连接图,'
          '会在基础策略上做增量探索,而不是从零重新学习。')

heading(doc, '3.6 关键调试经验', 2)
table(doc,
      ['现象', '根因', '修复'],
      [
          ['ECN 标记恒为 0', '单线程 BMv2 的 enq_qdepth 恒为 0,队列深度条件永不成立', '改用 deq_timedelta(排队等待时长)判拥塞'],
          ['ecn_ratio 读到但为 0', 'm.portIdx 1-based vs 寄存器 0-based 错位;标称除数 20000 偏大', 'portIdx 改 0-based;按实际发包量做标称'],
          ['BMv2 died', '上次残留 simple_switch 占 9090 端口 + IPC socket', 'cleanup() 改成 pkill 所有 simple_switch + 清 IPC'],
          ['UDP/TCP 不通', 'UDP 校验和不重算', 'h.udp.checksum = 0'],
          ['31 条流丢包 ~14-18%', 'Phase A 超发,单线程 BMv2 CPU 过载,交换机入口丢包', '节点约束:Phase A 速率 ×0.5,丢包降到 0%(阶段 3.7b)'],
      ],
      col_widths=[1.9, 2.4, 2.0])

heading(doc, '3.7 实验结果(实测)', 2)
heading(doc, '3.7a 首次自由通讯(改进前:丢包 14~18%)', 3)
code_block(doc, 'connection graph: 31 directional flows (16 nodes, 1~3 dests each)\n'
               'Phase A (超发): ecn_ratio=1.00 state=2\n'
               'Phase B (调度): active=31 mult=0.20\n'
               'sent: 39680  received: 34005  loss: 14.3%\n'
               'receiver InCsumErrors=0 RcvbufErrors=0(接收端约束全零)')
body(doc, '丢包原因(用数据定位):接收端校验/缓冲约束全为 0,排除接收端;交换机 ecn_marks ≈ egress_total'
          '(几乎每个包都在排队)→ 单线程 BMv2 处理速率约束被突破。31 条流在 Phase A 同时全速发送,'
          '总输入速率超过交换机处理上限,包在交换机入口被丢弃——属于"处理速率约束",非逻辑错误。')

heading(doc, '3.7b 节点约束优化(改进后:丢包 0%)', 3)
body(doc, '作为实验的一个改进阶段,对 Phase A 施加节点约束:把每流速率从 ×1.0 降到 ×0.5,'
          '使 31 条流的总负载落在单线程 BMv2 的处理能力内——这正是 RL 调度器 pacing 的思想(总速率控制在交换机能力内)。')
code_block(doc, 'Phase A (超发, mult=0.5): ecn_ratio=1.00 state=2   ← ECN 仍 100% 触发\n'
               'Phase B (调度): active=31 mult=0.20               ← DCQCN 量化降速不变\n'
               'sent: 39680  received: 39680  loss: 0.0%          ← 丢包从 14~18% 降到 0%')
body(doc, '保留的能力:ECN 标记(100%)、DCQCN 量化降速(×0.2)、0% 丢包——"拥塞检测 + 调度 + 无丢包"完整闭环达成。')

heading(doc, '3.8 MPTCP 多路径预留(已落地为实验四)', 2)
body(doc, 'flow.py 的 Flow/Subflow 抽象(按子流独立控速、可拆多条子流)为 MPTCP 预留了接口。'
          '该设计已在独立文件夹 mptcp_exp 中落地为完整实验(见下一章"实验四")。')

doc.add_page_break()

# ============ 实验四 ============
heading(doc, '4. 实验四:MPTCP 多路径 + DSN/SSN 二维流量控制(进行中)', 1)
body(doc, '在实验三基础上,把单条流升级为 MPTCP 风格的多子流连接:每对流随机生成 3~4 条子流'
          '(其中 1 条直连、不走交换机),用 DSN(数据序号)/SSN(子流序号)二维结构做流量控制,'
          '并让 DCQCN、Credit、RL 分域协同(见 MPTCP_DESIGN.md)。')

heading(doc, '4.1 文件与位置', 2)
body(doc, '实验文件夹:E:\\p4-workspace\\mptcp_exp(已同步到 GitHub mptcp_exp/)')
table(doc,
      ['文件', '作用'],
      [
          ['flow_mptcp.py', '数据模型:Flow(DSN序列+乱序缓冲) / Subflow(SSN序列+path=direct|sw);随机 3~4 子流其中 1 条直连'],
          ['run_mptcp.py', '主脚本:16命名空间+BMv2交换机+直连veth(专用/30子网绕过交换机);LPM路由;多里程碑演示'],
          ['mptcp_io.py', 'DSN/SSN 标记的 UDP 传输:发送端按 SSN 编号、接收端按 DSN 重组'],
          ['mptcp_scheduler.py', '三域拥塞控制:DCQCN(共享域)+ Credit(点对点)+ RL(全局)'],
          ['simple_router_global.p4', '数据面:转发 + ECN 标记(复用实验三)'],
          ['MPTCP_DESIGN.md', '设计蓝图:二维模型、三域架构、RL状态/动作/奖励、里程碑'],
      ],
      col_widths=[1.8, 4.2])

heading(doc, '4.2 核心概念:DSN/SSN 二维流量控制', 2)
body(doc, '每条 MPTCP 连接组织成二维结构:行方向 = 单条子流的 SSN 序列(子流内有序、可靠传输、credit);'
          '列方向 = 跨子流的 DSN 序列(连接级字节流,决定数据交给哪条子流、接收端重组)。')
code_block(doc, 'DSN 数据序号(连接级,唯一)\n'
               '  ├─ 子流1: SSN_1.0, SSN_1.1, ...   path=direct\n'
               '  ├─ 子流2: SSN_2.0, SSN_2.1, ...   path=sw\n'
               '  └─ 子流3: SSN_3.0, SSN_3.1, ...   path=sw\n'
               '接收端:按 DSN 把各子流数据重排成有序字节流')
body(doc, '流量控制同时作用于两个维度:SSN 维度(子流内限速/credit)与 DSN 维度(跨子流调度/路径选择)。')

heading(doc, '4.3 三域拥塞架构', 2)
table(doc,
      ['拥塞域', '拥塞性质', '机制', '信号', '动作'],
      [
          ['交换机(共享)', '所有经交换机子流共享瓶颈', 'DCQCN', 'ECN 标记', '经交换机子流量化降速'],
          ['直连(点对点)', '无共享', 'Credit', '接收方授信', '直连子流额度内发'],
          ['节点汇聚', '交换机子流+直连子流汇聚', 'Credit(按子流)', '各子流credit水位', '防独占缓冲'],
          ['全局', '跨所有子流', 'RL', 'ECN+credit+进度', '速率/路径/DSN分配'],
      ],
      col_widths=[1.4, 1.8, 1.1, 1.5, 1.9])

heading(doc, '4.4 已实现里程碑与实测', 2)
table(doc,
      ['里程碑', '内容', '实测结果'],
      [
          ['M1', 'DSN/SSN 数据模型 + 随机 3~4 子流(1直连)', '29 条流:29直连 + 79交换机子流'],
          ['M2', '直连 veth + 交换机双路径拓扑', '直连 ping 绕过交换机(rc=0)'],
          ['M3-M4', 'SSN 发送 + DSN 重组', '60段跨3子流全部按DSN有序重组(ordered=60,dup=0)'],
          ['M5-M7', 'DCQCN+Credit+RL 三域调度', 'ECN高时交换机子流降至0.25,直连保持0.75+'],
      ],
      col_widths=[1.0, 2.4, 2.9])
code_block(doc, '三域调度实测(ECN 波形):\n'
               '  ECN=0.8 state=2 → 交换机子流 avg_sw: 1.00→0.40(DCQCN 量化降速)\n'
               '  ECN=0.9 state=2 → 交换机子流 avg_sw →0.25(持续降速)\n'
               '  直连子流 avg_direct 始终 ~0.85+(不受交换机拥塞误伤)')

heading(doc, '4.5 待办里程碑', 2)
bullet(doc, 'M8:DSN/SSN 二维实时监控页(可视化每子流 SSN 进度 + 连接 DSN 重组)')
bullet(doc, '完整流量演示:多条流同时跑,RL 全局做速率/路径选择')
bullet(doc, '把真实 ECN 计数接入调度器(当前 demo 用合成 ECN 波形验证逻辑)')

doc.add_page_break()

# ============ 5. 阶段总结 ============
heading(doc, '5. 阶段总结(为论文准备)', 1)

heading(doc, '5.1 四个阶段一览', 2)
table(doc,
      ['阶段', '主题', '核心方法', '成果'],
      [
          ['阶段一', '双节点通讯', 'P4 IPv4 LPM 转发', '验证数据面转发通路,建立基础设施'],
          ['阶段二', '16节点 + 令牌桶限速', 'v1model meter 逐目的限速', '120/120 连通,限速 vs 不限速对比'],
          ['阶段三', '分层RL全局调度', '分层Q-learning + DCQCN + Credit', '自由双向通讯 + 0%接收端丢包'],
          ['阶段四', 'MPTCP多路径', 'DSN/SSN二维 + 三域拥塞控制', '多子流(直连+交换机)+ DSN重组 + 分域调度'],
      ],
      col_widths=[0.7, 1.7, 2.2, 2.7])

heading(doc, '5.2 关键技术贡献', 2)
bullet(doc, 'P4 数据面:转发 + ECN 标记(deq_timedelta 判拥塞),UDP 校验和清零')
bullet(doc, '软件侧分层 RL:高层选流 + 低层速率,交替冻结训练,指纹持久化 + 增量微调')
bullet(doc, 'DCQCN 拥塞控制:ECN 信号 → 量化降速 → 缓恢复')
bullet(doc, 'Credit 信用流控:接收方授信,防接收缓冲丢包')
bullet(doc, '丢包归因方法论:按位置(交换机入口/出口/接收校验/接收缓冲)分类约束')
bullet(doc, 'MPTCP:DSN/SSN 二维模型 + 直连/交换机双路径 + 三域拥塞控制')

heading(doc, '5.3 论文可展开方向', 2)
bullet(doc, '用 ns-3 做更真实的 RL 训练环境(替代简化流体模型)')
bullet(doc, 'MPTCP 路径选择:用流量控制指标(ECN/RTT/credit)做 DSN 到子流的智能分配')
bullet(doc, 'RL 算法升级:Q-learning → PPO/A2C + 连续状态')
bullet(doc, '多路径拓扑扩展:多交换机让每条子流真正独立路径')
bullet(doc, '把真实 ECN 计数接入 MPTCP 调度器闭环')

doc.save(r'E:\p4-workspace\实验记录_P4_BMv2.docx')
print('saved: E:\\p4-workspace\\实验记录_P4_BMv2.docx')
